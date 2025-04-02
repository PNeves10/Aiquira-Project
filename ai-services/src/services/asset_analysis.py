from typing import Dict, List, Optional
import openai
from pydantic import BaseModel
from datetime import datetime
import json
import PyPDF2
import io
import docx
from pathlib import Path

class FinancialMetrics(BaseModel):
    revenue: float
    profit: float
    employees: int
    year_founded: int

class Document(BaseModel):
    id: str
    name: str
    type: str
    url: str
    content: Optional[str] = None

class DocumentAnalysis(BaseModel):
    summary: str
    key_findings: List[str]
    risk_factors: List[Dict[str, str]]
    recommendations: List[str]
    financial_metrics: Optional[Dict[str, float]] = None
    legal_issues: Optional[List[str]] = None
    market_insights: Optional[List[str]] = None

class ValuationAnalysis(BaseModel):
    value: str
    confidence: float
    factors: List[str]

class RiskAnalysis(BaseModel):
    level: str  # 'low', 'medium', 'high'
    items: List[str]

class AIAnalysis(BaseModel):
    valuation: ValuationAnalysis
    risks: RiskAnalysis
    opportunities: List[str]
    market_trends: Optional[List[str]] = None
    competitor_analysis: Optional[Dict[str, List[str]]] = None

class AssetDetails(BaseModel):
    id: int
    title: str
    type: str
    price: str
    status: str
    description: str
    location: str
    financials: FinancialMetrics
    documents: List[Document]
    ai_analysis: Optional[AIAnalysis] = None

class AssetAnalysisService:
    def __init__(self, openai_api_key: str):
        self.openai_api_key = openai_api_key
        openai.api_key = openai_api_key

    def analyze_asset(self, asset: AssetDetails) -> AIAnalysis:
        """Perform AI analysis on an asset."""
        # Prepare the prompt for GPT-4
        prompt = self._create_analysis_prompt(asset)
        
        # Get analysis from OpenAI
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert M&A analyst specializing in business valuation and risk assessment."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        # Parse the response
        analysis = self._parse_analysis_response(response.choices[0].message.content)
        return analysis

    def _create_analysis_prompt(self, asset: AssetDetails) -> str:
        """Create a detailed prompt for the AI analysis."""
        return f"""
        Please analyze the following business asset and provide a comprehensive valuation and risk assessment:

        Asset Details:
        - Title: {asset.title}
        - Type: {asset.type}
        - Price: {asset.price}
        - Location: {asset.location}
        - Description: {asset.description}

        Financial Metrics:
        - Annual Revenue: {asset.financials.revenue}
        - Annual Profit: {asset.financials.profit}
        - Number of Employees: {asset.financials.employees}
        - Year Founded: {asset.financials.year_founded}

        Please provide:
        1. A detailed valuation analysis including:
           - Estimated value
           - Confidence level (0-100%)
           - Key factors affecting the valuation

        2. A risk assessment including:
           - Overall risk level (low/medium/high)
           - Specific risk factors

        3. Growth opportunities and potential

        4. Market trends analysis:
           - Current market conditions
           - Industry-specific trends
           - Future market predictions

        5. Competitor analysis:
           - Key competitors
           - Market positioning
           - Competitive advantages/disadvantages

        Format the response as a JSON object with the following structure:
        {{
            "valuation": {{
                "value": "string",
                "confidence": float,
                "factors": ["string"]
            }},
            "risks": {{
                "level": "string",
                "items": ["string"]
            }},
            "opportunities": ["string"],
            "market_trends": ["string"],
            "competitor_analysis": {{
                "competitors": ["string"],
                "market_position": ["string"],
                "advantages": ["string"],
                "disadvantages": ["string"]
            }}
        }}
        """

    def _parse_analysis_response(self, response: str) -> AIAnalysis:
        """Parse the AI response into structured data."""
        try:
            analysis_data = json.loads(response)
            return AIAnalysis(
                valuation=ValuationAnalysis(**analysis_data["valuation"]),
                risks=RiskAnalysis(**analysis_data["risks"]),
                opportunities=analysis_data["opportunities"],
                market_trends=analysis_data.get("market_trends"),
                competitor_analysis=analysis_data.get("competitor_analysis")
            )
        except (json.JSONDecodeError, KeyError) as e:
            raise ValueError(f"Failed to parse AI analysis response: {str(e)}")

    def analyze_documents(self, documents: List[Document]) -> List[DocumentAnalysis]:
        """Analyze documents related to the asset."""
        analyses = []
        
        for doc in documents:
            # Extract text content based on document type
            content = self._extract_document_content(doc)
            
            # Create analysis prompt
            prompt = self._create_document_analysis_prompt(doc, content)
            
            # Get analysis from OpenAI
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert in analyzing business documents for M&A transactions."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            
            # Parse the response
            analysis = self._parse_document_analysis_response(response.choices[0].message.content)
            analyses.append(analysis)
        
        return analyses

    def _extract_document_content(self, doc: Document) -> str:
        """Extract text content from different document types."""
        if not doc.content:
            return ""
            
        content = doc.content
        
        if doc.type.lower() == 'pdf':
            try:
                pdf_file = io.BytesIO(content.encode())
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                content = ""
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
            except Exception as e:
                print(f"Error extracting PDF content: {str(e)}")
                
        elif doc.type.lower() == 'docx':
            try:
                doc_file = io.BytesIO(content.encode())
                doc = docx.Document(doc_file)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as e:
                print(f"Error extracting DOCX content: {str(e)}")
                
        return content

    def _create_document_analysis_prompt(self, doc: Document, content: str) -> str:
        """Create a prompt for document analysis."""
        return f"""
        Please analyze the following {doc.type} document named "{doc.name}":

        Content:
        {content[:4000]}  # Limit content length for API

        Please provide:
        1. A concise summary of the document
        2. Key findings and important points
        3. Risk factors identified in the document
        4. Recommendations based on the document content
        5. Financial metrics if present
        6. Legal issues if any
        7. Market insights if available

        Format the response as a JSON object with the following structure:
        {{
            "summary": "string",
            "key_findings": ["string"],
            "risk_factors": [
                {{"category": "string", "severity": "string", "description": "string"}}
            ],
            "recommendations": ["string"],
            "financial_metrics": {{"metric_name": float}},
            "legal_issues": ["string"],
            "market_insights": ["string"]
        }}
        """

    def _parse_document_analysis_response(self, response: str) -> DocumentAnalysis:
        """Parse the document analysis response into structured data."""
        try:
            analysis_data = json.loads(response)
            return DocumentAnalysis(**analysis_data)
        except (json.JSONDecodeError, KeyError) as e:
            raise ValueError(f"Failed to parse document analysis response: {str(e)}")

    def generate_valuation_report(self, asset: AssetDetails, analysis: AIAnalysis) -> str:
        """Generate a detailed valuation report."""
        report = f"""
        Valuation Report for {asset.title}
        Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        
        1. Asset Overview
        ----------------
        Type: {asset.type}
        Location: {asset.location}
        Current Price: {asset.price}
        
        2. Financial Summary
        -------------------
        Annual Revenue: {asset.financials.revenue}
        Annual Profit: {asset.financials.profit}
        Employees: {asset.financials.employees}
        Year Founded: {asset.financials.year_founded}
        
        3. AI Valuation Analysis
        -----------------------
        Estimated Value: {analysis.valuation.value}
        Confidence Level: {analysis.valuation.confidence}%
        
        Key Factors:
        {chr(10).join(f"- {factor}" for factor in analysis.valuation.factors)}
        
        4. Risk Assessment
        ----------------
        Overall Risk Level: {analysis.risks.level.upper()}
        
        Risk Factors:
        {chr(10).join(f"- {risk}" for risk in analysis.risks.items)}
        
        5. Growth Opportunities
        ---------------------
        {chr(10).join(f"- {opportunity}" for opportunity in analysis.opportunities)}
        
        6. Market Analysis
        ----------------
        Market Trends:
        {chr(10).join(f"- {trend}" for trend in (analysis.market_trends or []))}
        
        Competitor Analysis:
        {chr(10).join(f"- {competitor}" for competitor in (analysis.competitor_analysis.get('competitors', []) if analysis.competitor_analysis else []))}
        
        Market Position:
        {chr(10).join(f"- {position}" for position in (analysis.competitor_analysis.get('market_position', []) if analysis.competitor_analysis else []))}
        
        Competitive Advantages:
        {chr(10).join(f"- {advantage}" for advantage in (analysis.competitor_analysis.get('advantages', []) if analysis.competitor_analysis else []))}
        
        Competitive Disadvantages:
        {chr(10).join(f"- {disadvantage}" for disadvantage in (analysis.competitor_analysis.get('disadvantages', []) if analysis.competitor_analysis else []))}
        """
        return report 